import json
import re
import sys
from pathlib import Path

from django.conf import settings
from docx import Document
from jsonschema import validate

"""
table heading row regular expressions used to identify and filter out table heading rows
"""
HEADINGS_LOOKUP = {
    0: "commodity|heading",
    1: "^description of product|^description of goods",
    2: "^working or processing|^conditions which must|^qualifying operation",
}

"""
Json schema to validate document structure before saving to file
"""
JSON_SCHEMA = {
    "$schema": "http://json-schema.org/draft-04/schema#",
    "type": "object",
    "properties": {
        "footnotes": {
            "type": "array",
            "items": [
                {
                    "type": "object",
                    "properties": {
                        "anchor": {"type": "string"},
                        "note": {"type": "string"},
                    },
                }
            ],
        },
        "chapters": {
            "type": "array",
            "items": [
                {
                    "type": "object",
                    "properties": {
                        "number": {"type": "string"},
                        "rows": {
                            "type": "array",
                            "items": [
                                {
                                    "type": "object",
                                    "properties": {
                                        "descriptions": {
                                            "type": "array",
                                            "items": [{"type": "string"}],
                                        },
                                        "ids": {
                                            "type": "array",
                                            "items": [{"type": "string"}],
                                        },
                                        "workingLeft": {
                                            "type": "array",
                                            "items": [{"type": "string"}],
                                        },
                                        "workingRight": {
                                            "type": "array",
                                            "items": [{"type": "string"}],
                                        },
                                    },
                                }
                            ],
                        },
                    },
                    "required": ["number", "rows"],
                }
            ],
        },
    },
    "required": ["chapters"],
}

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML

"""
Module that extract text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
PARA = WORD_NAMESPACE + "p"
TEXT = WORD_NAMESPACE + "t"


class DocxScraper:
    def __init__(self):
        self.table_dict = {"chapters": []}
        self.table_heading = None
        self.footnotes = []
        self.data_path = settings.OLD_RULES_OF_ORIGIN_DATA_PATH

    def load(self, docx_file):
        """
        Take a .docx file as an argument, read it from the filesystem and create an instance of a Document
        (from the python-docx library)
        extract the footnotes from the document
        loop through tables, filter out the duplicate cells, and process the data in each row to build a dictionary
        structure used to produce the json file output
        :param docx_file: file_name of the word document
        """

        document = Document(self.data_path.format(docx_file))
        # text = self.get_docx_text(self.data_path.format(docx_file))

        self.build_footnotes(document)
        heading_rows = []
        number_cols = []

        chapter_regex = re.compile(r"Chapter ([\d]+)")
        for table in document.tables:

            for row in table.rows:

                row = self.get_row(row)

                if self.is_table_heading(row):
                    heading_rows.append(row)
                if self.is_column_number_row(row):
                    number_cols.append(row)
                if self.is_chapter_row(row, chapter_regex):
                    chapter_number = chapter_regex.search(row[0][0]).group(1)
                    item = self.get_chapter_number(chapter_number)
                    item["rows"].append(self.create_rule_item(row))
                    self.table_dict["chapters"].append(item)

                else:
                    last_chap_idx = len(self.table_dict["chapters"]) - 1
                    if last_chap_idx != -1:

                        if len(row[0]) > 0:
                            if not self.is_table_heading(row):
                                self.table_dict["chapters"][last_chap_idx][
                                    "rows"
                                ].append(self.create_rule_item(row))
                        else:
                            last_row_index = (
                                len(self.table_dict["chapters"][last_chap_idx]["rows"])
                                - 1
                            )
                            keys = list(
                                self.table_dict["chapters"][last_chap_idx]["rows"][
                                    last_row_index
                                ].keys()
                            )

                            for idx, cell in enumerate(row):
                                if len(cell) > 0:
                                    for text in cell:
                                        self.table_dict["chapters"][last_chap_idx][
                                            "rows"
                                        ][last_row_index][keys[idx]].append(text)

        self.table_dict["footnotes"] = self.footnotes

        validate(instance=self.table_dict, schema=JSON_SCHEMA)

        # self.data_writer(self.data_path.format("import/{0}".format(docx_file+'_heading.json')), heading_rows)
        # self.data_writer(self.data_path.format("import/{0}".format(docx_file+'_columns.json')), number_cols)
        self.data_writer(
            self.data_path.format("import_v1/{0}".format(docx_file)), self.table_dict
        )

    @staticmethod
    def process_footnote(text):
        """
        Match footnote `(1)` in text and replace with in page anchor link
        :param text: text string
        :return: text string with html formatted footnote link
        """
        matched = re.match(u"\u2014", text)
        if matched:
            text = text.replace(u"\u2014", "- ")
        return re.sub(r"\(([\d]+)\)", '(<a href="#footnote_\\1">\\1</a>)', text)

    def create_rule_item(self, cells):
        """
        Check number of cells for to correctly match expected table output
        Pass each text string to process_footnotes method
        Assign text to array of cells used to build table rows for display

        :param cells: list of 3 or 4 lists of text
        :return: list of 4 lists of processed text with html footnote links
        """
        item = {}
        item_keys = ["id", "description", "workingLeft", "workingRight"]

        if len(cells) == 4:
            for i in range(len(item_keys)):
                item[item_keys[i]] = [self.process_footnote(text) for text in cells[i]]
        elif len(cells) == 3:
            cells.append([])
            for i in range(len(item_keys)):
                item[item_keys[i]] = [self.process_footnote(text) for text in cells[i]]
        else:
            for i in range(len(item_keys)):
                item[item_keys[i]] = [self.process_footnote(text) for text in cells[i]]

        return item

    @staticmethod
    def iter_unique_cells(row):
        """
        Generate cells in *row* skipping empty grid cells.
        :param row: docx table row
        :return: generator object of unique cells
        """

        prior_tc = None
        for cell in row.cells:
            this_tc = cell._tc
            if this_tc is prior_tc:
                continue
            prior_tc = this_tc
            yield cell

    @staticmethod
    def is_table_heading(row):
        """
        Check if the row is a heading row we do not need return True or False
        :param row: list is lists of text
        :return: Boolea
        """
        if len(row) > 3:
            return False
        matched_row = [
            [re.search(HEADINGS_LOOKUP[idx], text.lower()) for text in cell]
            for idx, cell in enumerate(row)
        ]
        match_tests = []
        for cell in matched_row:
            for match in cell:
                match_tests.append(True if match else False)
        return True in match_tests

    @staticmethod
    def get_chapter_number(number):
        """
        create a starter chapter dictionary object template with a number formatted as a zero padded number string
        :param number: string
        :return: dictionary object
        """
        return {"number": "{0:02d}".format(int(number)), "rows": []}

    def build_footnotes(self, document):
        """
        build a list of footnotes with html formatted link anchor and assign to the footnotes object property
        :param document: python-docx word docx object

        """
        pattern = r"^\(([\d]+)\)\s(.+)$"
        footnote_matches = [
            re.match(pattern, paragraph.text) for paragraph in document.paragraphs
        ]

        for match in footnote_matches:
            if match is not None:
                footnote = {
                    "anchor": '<a name="footnote_{0}">{0}</a>'.format(match.group(1)),
                    "note": match.group(2),
                }
                self.footnotes.append(footnote)

    def data_writer(self, file_path, data=None):
        """
        Write a file to the file system with data as content
        :param file_path: path to the file being created
        :param data: data to use as content for the file
        """
        if data is None:
            sys.exit()

        file_name = Path(file_path).stem.upper()

        outfile = open(
            self.data_path.format("archive/version1/import/{0}.json".format(file_name)),
            "w+",
        )
        json.dump(data, outfile)

    @staticmethod
    def is_column_number_row(row):
        """
        Check if the row is a column numbers row we do not need return True or False
        :param row: list is lists of text
        :return: Boolean
        """
        pattern = "^\(([\d]+)\)|or$"
        matched_row = [
            [re.search(pattern, text) for text in cell] for idx, cell in enumerate(row)
        ]
        match_tests = []
        for cell in matched_row:
            for match in cell:
                match_tests.append(True if match else False)
        return all(match is True for match in match_tests)

    @staticmethod
    def is_chapter_row(row, regex):
        """
        Check if the row is a chapter row we do need and return True or False
        :param row: list of lists of text
        :param regex: compiled regular expression
        :return: boolean
        """

        for idx, cell in enumerate(row):
            if idx < 1:
                for text in cell:
                    if regex.search(text) is not None:
                        return True
        return False

    def get_row(self, row):

        row_obj_1 = self.iter_unique_cells(row)
        row_obj_2 = self.iter_unique_cells(row)

        row1 = [
            [para.text for para in cell.paragraphs if para.text != ""]
            for cell in row_obj_1
        ]
        row2 = [
            [
                "".join(
                    [
                        "".join(
                            [
                                "".join([para.text for para in cell.paragraphs])
                                for cell in row.cells
                            ]
                        )
                        for row in table.rows
                    ]
                )
                for table in cell.tables
            ]
            for cell in row_obj_2
        ]

        for idx, item in enumerate(row1):
            if row2[idx]:
                for item in row2[idx]:
                    row1[idx].append(item)
        return row1
